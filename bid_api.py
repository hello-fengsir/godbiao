"""制标 API — 产品类别 / 产品线 / 技术参数 / 型号 / 预览 / Word导出"""
from fastapi import APIRouter, File, Form, UploadFile
from fastapi.responses import JSONResponse, Response
from io import BytesIO
from urllib.parse import quote

from models import (
    bid_categories_all, bid_categories_create, bid_categories_update, bid_categories_delete,
    bid_product_lines_all, bid_product_lines_create, bid_product_lines_update, bid_product_lines_delete,
    bid_parameters_all, bid_parameters_create, bid_parameters_update, bid_parameters_delete,
    bid_parameters_batch_create,
    bid_product_models_all, bid_product_models_create, bid_product_models_update, bid_product_models_delete,
)

router = APIRouter(prefix="/api/bid", tags=["制标"])


# ── 辅助函数 ──

def _require(data: dict, *fields):
    """校验必填字段，返回第一个缺失的字段名，全部通过返回 None"""
    for f in fields:
        if not data.get(f):
            return f
    return None


# ═══════════════════════════════════════════════════
# 类别
# ═══════════════════════════════════════════════════

@router.get("/categories")
def get_categories():
    return bid_categories_all()


@router.post("/categories")
def create_category(data: dict):
    miss = _require(data, "name")
    if miss:
        return JSONResponse({"error": f"缺少必填字段: {miss}"}, 400)
    return bid_categories_create(data["name"], data.get("sort_order", 0))


@router.put("/categories/{cat_id}")
def update_category(cat_id: int, data: dict):
    row = bid_categories_update(cat_id, data)
    if not row:
        return JSONResponse({"error": "类别不存在"}, 404)
    return dict(row)


@router.delete("/categories/{cat_id}")
def delete_category(cat_id: int):
    bid_categories_delete(cat_id)
    return {"ok": True}


# ═══════════════════════════════════════════════════
# 产品线
# ═══════════════════════════════════════════════════

@router.get("/product-lines")
def get_product_lines(category_id: int = 0):
    return bid_product_lines_all(category_id)


@router.post("/product-lines")
def create_product_line(data: dict):
    miss = _require(data, "category_id", "name")
    if miss:
        return JSONResponse({"error": f"缺少必填字段: {miss}"}, 400)
    return bid_product_lines_create(data["category_id"], data["name"], data.get("sort_order", 0))


@router.put("/product-lines/{line_id}")
def update_product_line(line_id: int, data: dict):
    row = bid_product_lines_update(line_id, data)
    if not row:
        return JSONResponse({"error": "产品线不存在"}, 404)
    return dict(row)


@router.delete("/product-lines/{line_id}")
def delete_product_line(line_id: int):
    bid_product_lines_delete(line_id)
    return {"ok": True}


# ═══════════════════════════════════════════════════
# 技术参数
# ═══════════════════════════════════════════════════

@router.get("/parameters")
def get_parameters(line_id: int = 0, param_type: str = ""):
    return bid_parameters_all(line_id, param_type)


@router.post("/parameters")
def create_parameter(data: dict):
    miss = _require(data, "line_id", "param_name")
    if miss:
        return JSONResponse({"error": f"缺少必填字段: {miss}"}, 400)
    return bid_parameters_create(
        data["line_id"], data["param_name"],
        data.get("param_type", "software"),
        data.get("param_value", ""),
        data.get("sort_order", 0),
    )


@router.put("/parameters/{param_id}")
def update_parameter(param_id: int, data: dict):
    row = bid_parameters_update(param_id, data)
    if not row:
        return JSONResponse({"error": "参数不存在"}, 404)
    return dict(row)


@router.delete("/parameters/{param_id}")
def delete_parameter(param_id: int):
    bid_parameters_delete(param_id)
    return {"ok": True}


@router.post("/parameters/upload")
async def upload_parameters(
    line_id: int = Form(...),
    param_type: str = Form("software"),
    file: UploadFile = File(...),
):
    fn = file.filename or ""
    if not fn.lower().endswith((".xlsx", ".xls")):
        return JSONResponse({"error": "仅支持 .xlsx / .xls 格式"}, 400)

    import openpyxl
    contents = await file.read()
    wb = openpyxl.load_workbook(BytesIO(contents), read_only=True, data_only=True)
    ws = wb.active

    rows_data = []
    for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
        if row_idx == 1:
            continue
        name = str(row[0]).strip() if row[0] else ""
        if not name:
            continue
        value = str(row[1]).strip() if len(row) > 1 and row[1] else ""
        rows_data.append((name, value))
    wb.close()

    imported = bid_parameters_batch_create(line_id, rows_data, param_type)
    return {"imported": imported}


# ═══════════════════════════════════════════════════
# 产品型号
# ═══════════════════════════════════════════════════

@router.get("/product-models")
def get_product_models(line_id: int = 0):
    return bid_product_models_all(line_id)


@router.post("/product-models")
def create_product_model(data: dict):
    miss = _require(data, "line_id", "name")
    if miss:
        return JSONResponse({"error": f"缺少必填字段: {miss}"}, 400)
    return bid_product_models_create(
        data["line_id"], data["name"],
        data.get("description", ""),
        data.get("sort_order", 0),
    )


@router.put("/product-models/{model_id}")
def update_product_model(model_id: int, data: dict):
    row = bid_product_models_update(model_id, data)
    if not row:
        return JSONResponse({"error": "型号不存在"}, 404)
    return dict(row)


@router.delete("/product-models/{model_id}")
def delete_product_model(model_id: int):
    bid_product_models_delete(model_id)
    return {"ok": True}


# ═══════════════════════════════════════════════════
# 预览 & Word 导出
# ═══════════════════════════════════════════════════

@router.post("/generate/preview")
def generate_preview(data: dict):
    line_id = data.get("line_id", 0)
    param_ids = data.get("param_ids", [])

    from models import get_db
    with get_db() as db:
        line = db.execute(
            "SELECT bpl.*, bc.name as cat_name FROM bid_product_lines bpl "
            "JOIN bid_categories bc ON bc.id=bpl.category_id WHERE bpl.id=?",
            (line_id,)
        ).fetchone()
        if not line:
            return JSONResponse({"error": "产品线不存在"}, 404)

        if param_ids:
            placeholders = ",".join("?" * len(param_ids))
            rows = db.execute(
                f"SELECT * FROM bid_parameters WHERE id IN ({placeholders}) ORDER BY param_type DESC, sort_order, id",
                param_ids,
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM bid_parameters WHERE line_id=? ORDER BY param_type DESC, sort_order, id",
                (line_id,),
            ).fetchall()

    return {
        "category": line["cat_name"],
        "product_line": line["name"],
        "parameters": [
            {"id": r["id"], "name": r["param_name"], "value": r["param_value"], "type": r["param_type"]}
            for r in rows
        ],
        "total": len(rows),
    }


@router.post("/generate/docx")
async def generate_docx(data: dict):
    line_id = data.get("line_id", 0)
    param_ids = data.get("param_ids", [])

    from models import get_db
    with get_db() as db:
        line = db.execute(
            "SELECT bpl.*, bc.name as cat_name FROM bid_product_lines bpl "
            "JOIN bid_categories bc ON bc.id=bpl.category_id WHERE bpl.id=?",
            (line_id,)
        ).fetchone()
        if not line:
            return JSONResponse({"error": "产品线不存在"}, 404)

        if param_ids:
            placeholders = ",".join("?" * len(param_ids))
            rows = db.execute(
                f"SELECT * FROM bid_parameters WHERE id IN ({placeholders}) ORDER BY param_type DESC, sort_order, id",
                param_ids,
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM bid_parameters WHERE line_id=? ORDER BY param_type DESC, sort_order, id",
                (line_id,),
            ).fetchall()

    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    title = doc.add_heading("技术参数表", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{line['cat_name']} — {line['name']}")
    run.font.size = Pt(14)

    doc.add_paragraph()

    cat_pattern = re.compile(r'^\[(.+?)\]\s*(.*)')
    grouped = {}
    group_order = []
    for r in rows:
        m = cat_pattern.match(r['param_name'] or '')
        if m:
            cat = m.group(1).strip()
            clean = m.group(2).strip()
        else:
            cat = '其他'
            clean = (r['param_name'] or '').strip()
        if cat not in grouped:
            grouped[cat] = []
            group_order.append(cat)
        grouped[cat].append(clean)

    global_idx = 0
    for cat in group_order:
        items = grouped[cat]
        doc.add_heading(cat, level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "技术参数名称"

        for clean_name in items:
            global_idx += 1
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(global_idx)
            cells[1].text = clean_name

        doc.add_paragraph()

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"共 {len(rows)} 项参数")
    run.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = line['name'].replace("/", "_").replace(" ", "_")
    filename = f"技术参数_{safe_name}.docx"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )
