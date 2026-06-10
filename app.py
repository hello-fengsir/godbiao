"""
神奇阿标 (God Biao) 主应用
"""
import json
import asyncio
import uuid
import sqlite3
import os
import shutil
import glob
from datetime import datetime, timedelta

from fastapi import FastAPI, File, Form, UploadFile, Request
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from config import BASE_DIR, UPLOAD_DIR, DB_PATH, MODEL_PRESETS
from parser import parse_file
from llm_client import compare_bid, evaluate_format, extract_key_info, test_api_key
from models import get_db, init_product_tables

from starlette.templating import Jinja2Templates as StarletteJinja2

app = FastAPI(title="神奇阿标 - God Biao")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
templates = StarletteJinja2(directory=str(BASE_DIR / "templates"))

UPLOAD_DIR.mkdir(exist_ok=True)

# 比对进度跟踪（内存中）
_compare_progress: dict = {}  # job_id → {phase, content_done, format_done, started_at, error}
MAX_HISTORY = 3


def init_db():
    with sqlite3.connect(str(DB_PATH)) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                id TEXT PRIMARY KEY,
                created_at TEXT,
                req_text TEXT,
                bid_filename TEXT,
                bid_text TEXT,
                result_json TEXT,
                tokens_used INTEGER DEFAULT 0,
                cost_estimate REAL DEFAULT 0,
                model_used TEXT DEFAULT '',
                eval_mode TEXT DEFAULT 'combined'
            )
        """)
        # 兼容旧表：尝试添加 eval_mode 列
        try:
            conn.execute("ALTER TABLE jobs ADD COLUMN eval_mode TEXT DEFAULT 'combined'")
        except:
            pass
        conn.commit()


def cleanup_on_refresh():
    """每次刷新首页时清理：上传文件 + 超过3天的旧任务"""
    # 1. 清理上传目录中的所有文件
    for f in glob.glob(str(UPLOAD_DIR / "*")):
        try:
            os.remove(f) if os.path.isfile(f) else shutil.rmtree(f)
        except Exception:
            pass

    # 2. 删除超过3天的数据库记录
    cutoff = (datetime.now() - timedelta(days=1)).isoformat()
    try:
        with sqlite3.connect(str(DB_PATH)) as conn:
            conn.execute("DELETE FROM jobs WHERE created_at < ?", (cutoff,))
            conn.commit()
    except Exception:
        pass

    # 3. 清理进度缓存
    _compare_progress.clear()


init_db()


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    cleanup_on_refresh()
    tpl = templates.env.get_template("index.html")
    return HTMLResponse(tpl.render({"request": request}))


@app.post("/upload")
async def upload(
    request: Request,
    bid_file: UploadFile = File(None),
    tech_file: UploadFile = File(None),
    biz_file: UploadFile = File(None),
    requirements: str = Form(""),
    req_file: UploadFile = File(None),
    eval_mode: str = Form("combined"),
):
    bid_text = ""
    bid_filename = ""

    if eval_mode == "separate":
        # 分离模式：技术标 + 商务标
        tech_text = ""
        biz_text = ""
        if tech_file and tech_file.filename:
            tech_text = await parse_file(tech_file)
        if biz_file and biz_file.filename:
            biz_text = await parse_file(biz_file)
        if not tech_text and not biz_text:
            return templates.TemplateResponse(
                request,
                "index.html",
                {"error": "请至少上传技术标或商务标文件"},
            )
        bid_text = ""
        if tech_text:
            bid_text += "=== 技术标 ===\n" + tech_text + "\n\n"
        if biz_text:
            bid_text += "=== 商务标 ===\n" + biz_text
        bid_filename = (tech_file.filename if tech_file else "") + (" + " + biz_file.filename if biz_file else "")
        print(f"upload (separate): tech={tech_file.filename if tech_file else 'N/A'}, biz={biz_file.filename if biz_file else 'N/A'}")
    else:
        # 合并模式：单文件
        if not bid_file or not bid_file.filename:
            return templates.TemplateResponse(
                request,
                "index.html",
                {"error": "请上传投标文件"},
            )
        bid_text = await parse_file(bid_file)
        bid_filename = bid_file.filename
        print(f"upload: {bid_file.filename}")

    req_text = requirements.strip()
    if req_file:
        req_text += "\n" + (await parse_file(req_file))

    if not bid_text:
        return templates.TemplateResponse(
            request,
            "index.html",
            {"error": "无法解析投标文件，请检查文件格式（支持 PDF/Word/图片）"},
        )

    job_id = str(uuid.uuid4())[:8]
    with sqlite3.connect(str(DB_PATH)) as conn:
        conn.execute(
            "INSERT INTO jobs (id, created_at, req_text, bid_filename, bid_text, eval_mode) VALUES (?,?,?,?,?,?)",
            (job_id, datetime.now().isoformat(), req_text, bid_filename, bid_text, eval_mode),
        )
        conn.commit()
        # 保留最多 MAX_HISTORY 条记录
        conn.execute("DELETE FROM jobs WHERE id NOT IN (SELECT id FROM jobs ORDER BY created_at DESC LIMIT ?)", (MAX_HISTORY,))
        conn.commit()

    tpl = templates.env.get_template("review.html")
    return HTMLResponse(tpl.render({
        "request": request,
        "job_id": job_id,
        "req_text": req_text,
        "bid_text_preview": bid_text,
        "bid_filename": bid_filename,
        "bid_full_len": len(bid_text),
        "req_full_len": len(req_text),
        "eval_mode": eval_mode,
    }))


@app.get("/api/model-presets")
async def get_model_presets():
    """返回预置模型列表（不含任何 API Key）"""
    return {"presets": MODEL_PRESETS}


@app.get("/api/jobs")
async def list_jobs():
    """列出历史比对记录"""
    with sqlite3.connect(str(DB_PATH)) as conn:
        rows = conn.execute(
            "SELECT id, created_at, bid_filename, model_used, tokens_used FROM jobs ORDER BY created_at DESC"
        ).fetchall()
    return {
        "jobs": [
            {
                "id": r[0],
                "created_at": r[1],
                "bid_filename": r[2],
                "model_used": r[3],
                "tokens_used": r[4],
            }
            for r in rows
        ]
    }


@app.delete("/api/jobs/{job_id}")
async def delete_job(job_id: str):
    """删除指定比对记录"""
    with sqlite3.connect(str(DB_PATH)) as conn:
        conn.execute("DELETE FROM jobs WHERE id=?", (job_id,))
        conn.commit()
    return {"ok": True}


@app.post("/api/test-key")
async def test_key(request: Request):
    """测试 API Key 是否有效"""
    body = await request.json()
    base_url = body.get("base_url", "")
    api_key = body.get("api_key", "")
    model = body.get("model", "")
    if not base_url or not api_key:
        return JSONResponse({"ok": False, "detail": "缺少 base_url 或 api_key"}, status_code=400)
    result = await test_api_key(base_url, api_key, model)
    return result


@app.post("/compare/{job_id}")
async def run_comparison(job_id: str, request: Request):
    # 从请求头获取 API 配置
    api_key = request.headers.get("X-API-Key", "")
    base_url = request.headers.get("X-API-Base-Url", "")
    model = request.headers.get("X-API-Model", "")
    provider_name = request.headers.get("X-API-Provider-Name", "")
    
    provider = None
    if api_key and base_url and model:
        provider = {
            "name": provider_name,
            "base_url": base_url,
            "api_key": api_key,
            "model": model,
        }

    with sqlite3.connect(str(DB_PATH)) as conn:
        row = conn.execute(
            "SELECT req_text, bid_text, eval_mode FROM jobs WHERE id=?", (job_id,)
        ).fetchone()
        if not row:
            return JSONResponse({"error": "任务不存在"}, status_code=404)

        req_text, bid_text, job_eval_mode = row[0], row[1], (row[2] if len(row) > 2 else "combined")

    # 初始化进度
    started = datetime.now().isoformat()
    _compare_progress[job_id] = {
        "phase": "starting", "content_done": False, "format_done": False,
        "started_at": started, "error": None,
    }

    try:
        # 包装任务：完成后更新进度
        async def run_content():
            try:
                r = await compare_bid(req_text, bid_text, provider, job_eval_mode)
                _compare_progress[job_id]["content_done"] = True
                _compare_progress[job_id]["phase"] = "format"
                return r
            except Exception as e:
                _compare_progress[job_id]["error"] = f"内容比对: {e}"
                raise

        async def run_format():
            try:
                r = await evaluate_format(bid_text, provider)
                _compare_progress[job_id]["format_done"] = True
                _compare_progress[job_id]["phase"] = "saving"
                return r
            except Exception as e:
                _compare_progress[job_id]["format_done"] = True
                _compare_progress[job_id]["phase"] = "saving"
                print(f"[Format] 格式评估失败（非致命）: {e}")
                return {"overall": "格式评估暂不可用", "score": 0, "items": [], "tokens_used": 0}

        async def run_key_info():
            try:
                r = await extract_key_info(bid_text, provider)
                return r
            except Exception as e:
                print(f"[KeyInfo] 关键信息提取失败（非致命）: {e}")
                return {"business": {}, "technical": {}, "tokens_used": 0}

        content_result, format_result, key_info = await asyncio.gather(
            run_content(), run_format(), run_key_info(), return_exceptions=True
        )

        # 处理内容比对异常
        if isinstance(content_result, Exception):
            _compare_progress[job_id]["phase"] = "error"
            _compare_progress[job_id]["error"] = str(content_result)
            return JSONResponse({"error": f"内容比对失败: {str(content_result)}"}, status_code=500)
    except Exception as e:
        import traceback
        traceback.print_exc()
        _compare_progress[job_id]["phase"] = "error"
        _compare_progress[job_id]["error"] = str(e)
        return JSONResponse({"error": f"比对失败: {str(e)}"}, status_code=500)

    content_items = content_result.get("items", [])
    content_overall = content_result.get("overall", "")
    tokens = (content_result.get("tokens_used", 0) +
              format_result.get("tokens_used", 0) +
              (key_info.get("tokens_used", 0) if isinstance(key_info, dict) else 0))
    model_name = content_result.get("model", "")
    cost = round(tokens * 0.000002, 6)

    # 处理 key_info 异常
    if isinstance(key_info, Exception) or not isinstance(key_info, dict):
        key_info = {"business": {}, "technical": {}}

    # 合并存储: content + format + key_info
    combined_result = {
        "content": {"items": content_items, "overall": content_overall},
        "format": format_result,
        "key_info": key_info,
    }

    with sqlite3.connect(str(DB_PATH)) as conn:
        conn.execute(
            "UPDATE jobs SET result_json=?, tokens_used=?, cost_estimate=?, model_used=? WHERE id=?",
            (json.dumps(combined_result, ensure_ascii=False), tokens, cost, model_name, job_id),
        )
        conn.commit()

    _compare_progress[job_id]["phase"] = "done"

    # 计算格式评分
    fmt_score = format_result.get("score", 0)
    fmt_overall = format_result.get("overall", "")
    fmt_items = format_result.get("items", [])

    return {
        "job_id": job_id,
        "content": {"items": content_items, "overall": content_overall},
        "format": {"score": fmt_score, "overall": fmt_overall, "items": fmt_items},
        "key_info": key_info,
        "tokens": tokens, "cost": cost, "model": model_name,
    }


@app.get("/api/compare-status/{job_id}")
async def compare_status(job_id: str):
    """获取比对进度"""
    p = _compare_progress.get(job_id)
    if not p:
        return {"phase": "unknown", "message": "任务未找到"}
    return {
        "phase": p["phase"],
        "content_done": p["content_done"],
        "format_done": p["format_done"],
        "started_at": p["started_at"],
        "error": p.get("error"),
    }


@app.get("/result/{job_id}", response_class=HTMLResponse)
async def result_page(request: Request, job_id: str):
    with sqlite3.connect(str(DB_PATH)) as conn:
        row = conn.execute(
            "SELECT req_text, bid_filename, result_json, tokens_used, model_used, eval_mode FROM jobs WHERE id=?",
            (job_id,),
        ).fetchone()

    if not row:
        tpl = templates.env.get_template("index.html")
        return HTMLResponse(tpl.render({"request": request, "error": "任务不存在"}))

    results = json.loads(row[2]) if row[2] else {}
    overall = ""
    format_result = {}
    key_info = {"business": {}, "technical": {}}
    content_items = []

    if isinstance(results, dict):
        # 新格式: {"content": {...}, "format": {...}, "key_info": {...}}
        if "content" in results:
            content_data = results["content"]
            content_items = content_data.get("items", [])
            overall = content_data.get("overall", "")
            format_result = results.get("format", {})
            key_info = results.get("key_info", {"business": {}, "technical": {}})
        # 旧格式兼容: {"items": [...], "overall": "..."}
        elif "items" in results:
            overall = results.get("overall", "")
            content_items = results.get("items", [])
        else:
            content_items = []
    elif isinstance(results, list):
        content_items = results

    eval_mode_display = row[5] if len(row) > 5 else "combined"
    total = len(content_items) or 1
    full_pass = sum(1 for r in content_items if r.get("status") == "满足")
    partial = sum(1 for r in content_items if r.get("status") == "部分满足")
    failed = sum(1 for r in content_items if r.get("status") in ("不满足", "缺失"))

    # 格式评估数据
    fmt_score = format_result.get("score", 0) if format_result else 0
    fmt_overall = format_result.get("overall", "") if format_result else ""
    fmt_items = format_result.get("items", []) if format_result else []

    tpl = templates.env.get_template("result.html")
    return HTMLResponse(tpl.render({
            "request": request,
            "job_id": job_id,
            "req_text": row[0],
            "bid_filename": row[1],
            "results": content_items,
            "tokens": row[3],
            "model": row[4],
            "total": len(content_items),
            "full_pass": full_pass,
            "full_pass_pct": round(full_pass / total * 100, 1),
            "partial": partial,
            "partial_pct": round(partial / total * 100, 1),
            "failed": failed,
            "failed_pct": round(failed / total * 100, 1),
            "eval_mode": eval_mode_display,
            "overall": overall,
            "fmt_score": fmt_score,
            "fmt_overall": fmt_overall,
            "fmt_items": fmt_items,
            "key_info": key_info,
        }))


# ── 制标模块 API ──

from io import BytesIO
from collections import defaultdict

@app.get("/bidding", response_class=HTMLResponse)
def bidding_page(request: Request):
    return templates.TemplateResponse(request, "bidding.html", {})

@app.get("/api/categories")
def api_categories():
    """返回所有产品分类"""
    with get_db() as db:
        init_product_tables()
        rows = db.execute("SELECT DISTINCT category FROM products ORDER BY category").fetchall()
        return {"categories": [r["category"] for r in rows]}

@app.get("/api/products")
def api_products(category: str = ""):
    """返回指定分类下的产品列表"""
    with get_db() as db:
        init_product_tables()
        if category:
            rows = db.execute(
                "SELECT id, category, name, model, specs, notes FROM products WHERE category=? ORDER BY name, model",
                (category,)
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT id, category, name, model, specs, notes FROM products ORDER BY category, name, model"
            ).fetchall()
        return {"products": [dict(r) for r in rows]}

@app.post("/api/export-word")
async def export_word(data: dict):
    """导出选中产品为 Word 表格，按模块标签分组"""
    items = data.get("items", [])
    if not items:
        return JSONResponse({"error": "未选择产品"}, status_code=400)

    product_ids = [it["product_id"] for it in items]
    module_map = {it["product_id"]: it.get("module_tag", "未分类") for it in items}

    with get_db() as db:
        placeholders = ",".join("?" * len(product_ids))
        rows = db.execute(
            f"SELECT * FROM products WHERE id IN ({placeholders}) ORDER BY category, name, model",
            product_ids
        ).fetchall()

    try:
        from docx import Document
    except ImportError:
        return JSONResponse({"error": "python-docx 未安装"}, status_code=500)

    doc = Document()
    doc.add_heading("投标产品参数表", level=1)
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    groups = defaultdict(lambda: defaultdict(list))
    for r in rows:
        module = module_map.get(r["id"], "未分类")
        groups[module][r["category"]].append(r)

    for module, cats in groups.items():
        doc.add_heading(f"📦 {module}", level=2)
        for cat, cat_items in cats.items():
            doc.add_heading(cat, level=3)
            table = doc.add_table(rows=1, cols=4, style="Table Grid")
            hdr = table.rows[0].cells
            hdr[0].text = "产品名"
            hdr[1].text = "型号"
            hdr[2].text = "参数"
            hdr[3].text = "备注"
            for item in cat_items:
                row_cells = table.add_row().cells
                row_cells[0].text = item["name"]
                row_cells[1].text = item["model"]
                row_cells[2].text = item["specs"]
                row_cells[3].text = item["notes"] or ""

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    from urllib.parse import quote
    filename = quote("投标参数表.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8880)