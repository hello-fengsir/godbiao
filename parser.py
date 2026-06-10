"""
文件解析模块 - 支持 PDF / Word / 图片(OCR)
增强版：提取表格 + 内嵌图片 OCR
"""
import io
import os
import tempfile
from fastapi import UploadFile
from PIL import Image


async def parse_file(file: UploadFile) -> str:
    """解析上传文件，返回纯文本（含表格和图片OCR结果）"""
    content = await file.read()
    filename = (file.filename or "").lower()
    await file.seek(0)

    if filename.endswith(".pdf"):
        return parse_pdf(content)
    elif filename.endswith((".docx", ".doc")):
        return parse_docx(content)
    elif filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif")):
        return parse_image(content)
    elif filename.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    else:
        return content.decode("utf-8", errors="replace")


# ─────────────────────── PDF 解析 ───────────────────────

def parse_pdf(content: bytes) -> str:
    parts = []

    # 1. 文本提取 (pdfplumber，包含表格)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                # 提取文本
                txt = page.extract_text()
                if txt:
                    parts.append(f"\n=== 第{i}页 ===\n{txt}")

                # 提取表格
                tables = page.extract_tables()
                if tables:
                    for ti, table in enumerate(tables):
                        parts.append(format_table(table, f"第{i}页 表格{ti+1}"))
    except Exception as e:
        parts.append(f"[PDF文本解析失败: {e}]")

    # 2. 内嵌图片提取 + OCR (pymupdf)
    try:
        img_texts = extract_pdf_images(content)
        if img_texts:
            parts.append(img_texts)
    except Exception as e:
        parts.append(f"\n[PDF图片提取失败: {e}]")

    return "\n".join(parts)


def extract_pdf_images(content: bytes) -> str:
    """从 PDF 中提取内嵌图片并 OCR"""
    import fitz  # pymupdf

    texts = []
    doc = fitz.open(stream=content, filetype="pdf")

    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)

        for img_idx, img_info in enumerate(image_list):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                img = Image.open(io.BytesIO(image_bytes))
                ocr_text = ocr_image(img)

                if ocr_text and len(ocr_text.strip()) > 5:
                    texts.append(f"\n[PDF第{page_num+1}页内嵌图片{img_idx+1} OCR结果]\n{ocr_text}")
            except Exception:
                pass

    doc.close()
    return "\n".join(texts) if texts else ""


# ─────────────────────── Word 解析 ───────────────────────

def parse_docx(content: bytes) -> str:
    parts = []

    try:
        from docx import Document
        doc = Document(io.BytesIO(content))

        # 段落文本
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        if para_texts:
            parts.append("\n".join(para_texts))

        # 表格提取
        if doc.tables:
            parts.append("\n=== 文档中的表格 ===")
            for ti, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                parts.append(format_table(rows, f"表格{ti+1}"))

        # 内嵌图片提取 + OCR
        img_texts = extract_docx_images(content)
        if img_texts:
            parts.append(img_texts)

    except Exception as e:
        parts.append(f"[Word解析失败: {e}]")

    return "\n".join(parts)


def extract_docx_images(content: bytes) -> str:
    """从 docx 中提取内嵌图片并 OCR"""
    import zipfile
    from io import BytesIO

    texts = []
    try:
        with zipfile.ZipFile(BytesIO(content)) as z:
            image_names = [n for n in z.namelist() if n.startswith("word/media/") and
                          n.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.tif', '.webp'))]
            for img_name in image_names:
                try:
                    img_data = z.read(img_name)
                    img = Image.open(BytesIO(img_data))
                    ocr_text = ocr_image(img)
                    if ocr_text and len(ocr_text.strip()) > 5:
                        label = os.path.basename(img_name)
                        texts.append(f"\n[Word内嵌图片 {label} OCR结果]\n{ocr_text}")
                except Exception:
                    pass
    except Exception:
        pass

    return "\n".join(texts) if texts else ""


# ─────────────────────── 图片 OCR ───────────────────────

def parse_image(content: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(content))
        text = ocr_image(img)
        return text.strip() or "[OCR未识别到文字]"
    except Exception as e:
        return f"[OCR失败: {e}]"


def ocr_image(img: Image.Image) -> str:
    """对 PIL Image 执行 OCR"""
    try:
        import pytesseract
        # 预处理：放大太小文字
        w, h = img.size
        if w < 300 or h < 300:
            img = img.resize((max(w, 300), max(h, 300)), Image.Resampling.LANCZOS)

        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return text.strip()
    except Exception as e:
        return f"[OCR失败: {e}]"


# ─────────────────────── 表格格式化 ───────────────────────

def format_table(rows: list, title: str = "") -> str:
    """将二维列表格式化为易读的文本表格"""
    if not rows:
        return ""

    # 计算每列最大宽度
    col_widths = []
    for row in rows:
        for ci, cell in enumerate(row):
            cell_len = len(str(cell or ""))
            if ci >= len(col_widths):
                col_widths.append(cell_len)
            else:
                col_widths[ci] = max(col_widths[ci], cell_len)

    # 限制列宽避免过宽
    col_widths = [min(w, 40) for w in col_widths]

    def fmt_cell(text, width):
        t = str(text or "").replace("\n", " ")
        if len(t) > width:
            t = t[:width-2] + ".."
        return t.ljust(width)

    lines = [f"\n### {title}" if title else ""]
    # 表头
    if rows:
        header = " | ".join(fmt_cell(c, col_widths[i]) for i, c in enumerate(rows[0]))
        sep = "-+-".join("-" * col_widths[i] for i in range(len(col_widths)))
        lines.append(header)
        lines.append(sep)
        for row in rows[1:]:
            line = " | ".join(fmt_cell(c, col_widths[i]) for i, c in enumerate(row))
            lines.append(line)

    return "\n".join(lines)


# ─────────────────────── 目录检测 ───────────────────────

def detect_toc(text: str) -> str:
    """扫描文本前 30%，检测目录结构。返回提示标记或空字符串。"""
    import re
    lines = text.split("\n")
    # 只看前 30% 的行
    sample = lines[:max(10, int(len(lines) * 0.3))]

    toc_lines = []
    toc_start = -1
    in_toc = False
    seen_toc_title = False

    # 目录标题关键词
    toc_titles = ["目录", "目次", "目  录", "CONTENTS", "Table of Contents", "目    录"]

    for i, line in enumerate(sample):
        stripped = line.strip()
        if not stripped:
            if in_toc and len(toc_lines) >= 3:
                continue  # 目录内空行允许
            else:
                in_toc = False
                continue

        # 检测目录标题
        if not seen_toc_title:
            for t in toc_titles:
                if t.lower() in stripped.lower() and len(stripped) < 30:
                    seen_toc_title = True
                    toc_start = i
                    in_toc = True
                    break

        if in_toc:
            # 目录行特征：末尾有数字（页码）
            # 模式：...数字 或 空格+数字
            has_page_num = bool(re.search(r'\.{2,}\s*\d+$|[\s.]\d{1,4}$|\d{1,4}$', stripped))
            # 排除纯数字行或太短的行
            is_not_page_num_only = not re.match(r'^\d{1,4}$', stripped)
            # 有层级编号
            has_numbering = bool(re.search(r'^[\d一二三四五六七八九十]+[\.\、\s]|^第[一二三四五六七八九十\d]+[章节]', stripped))
            
            if has_page_num and is_not_page_num_only:
                toc_lines.append(i)
            elif has_numbering and len(stripped) > 4:
                # 也可能没有页码但仍是目录行
                toc_lines.append(i)
            elif len(toc_lines) == 0:
                # 刚开始，还没匹配到
                continue
            elif has_page_num:
                toc_lines.append(i)

    # 判定：至少 3 行目录条目
    if seen_toc_title or len(toc_lines) >= 5:
        first = toc_lines[0] if toc_lines else toc_start
        last = toc_lines[-1] if toc_lines else toc_start + 1
        return f"\n[✅ 系统检测到目录结构：第{first+1}-{last+1}行疑似为目录页，共{len(toc_lines) or '多'}个章节条目]\n"

    # 如果没有明确的"目录"标题，但有很多疑似目录行
    if len(toc_lines) >= 8:
        first, last = toc_lines[0], toc_lines[-1]
        return f"\n[✅ 系统检测到疑似目录结构：第{first+1}-{last+1}行含{len(toc_lines)}个疑似章节条目]\n"

    return ""